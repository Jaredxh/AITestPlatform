"""需求文档解析器：支持 Word (.docx / .doc) 和 PDF 格式的文本提取。"""

import io
import logging
import os
import quopri
import re
import subprocess
import tempfile
from email import policy
from email.parser import BytesParser
from html import unescape

logger = logging.getLogger(__name__)

SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",  # .doc
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

DOC_EXT_HINTS = {"doc", "docx", "pdf"}


def validate_file(content_type: str, size: int, filename: str) -> None:
    if size > MAX_FILE_SIZE:
        raise ValueError(f"文件大小超过限制（最大 {MAX_FILE_SIZE // 1024 // 1024}MB）")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if content_type in SUPPORTED_CONTENT_TYPES or ext in DOC_EXT_HINTS:
        return

    raise ValueError(
        f"不支持的文件类型: {content_type}，仅支持 .pdf / .docx / .doc"
    )


def extract_text(content: bytes, content_type: str, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if content_type == "application/pdf" or ext == "pdf":
        return _extract_pdf(content)
    if (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == "docx"
    ):
        return _extract_docx(content)
    if content_type == "application/msword" or ext == "doc":
        return _extract_doc(content)

    raise ValueError(f"无法解析文件类型: {content_type}")


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"--- 第 {i + 1} 页 ---\n{text.strip()}")
    if not pages:
        raise ValueError("PDF 文件无法提取到文本内容，可能是扫描件或图片型 PDF")
    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        raise ValueError("Word 文档中未提取到文本内容")
    return "\n\n".join(parts)


def _extract_doc(content: bytes) -> str:
    """老版 .doc (OLE compound) 文本抽取。

    流水线（任一成功即返回）：
      1. antiword  -m UTF-8.txt   首选，国内 Word 字体覆盖较好
      2. antiword                 不带映射兜底
      3. catdoc    -d utf-8       antiword 解析不出来时的备选
      4. 直接尝试当作 docx 读取    用户把 .docx 错命名成 .doc 时也能救
      5. 最后才报错               并附带 antiword/catdoc 的 stderr，便于排查

    任意一步异常都不抛，只标记"未能解析"，由调用方决定是降级保存还是 422。
    """
    html_text = _extract_pseudo_doc_html(content)
    if html_text:
        return html_text

    # 0) 误命名为 .doc 的 docx：先用 zip 嗅探
    if content[:4] == b"PK\x03\x04":
        try:
            return _extract_docx(content)
        except Exception:  # noqa: BLE001
            pass

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as fh:
        fh.write(content)
        tmp_path = fh.name

    errors: list[str] = []
    try:
        for cmd in (
            ["antiword", "-m", "UTF-8.txt", tmp_path],
            ["antiword", tmp_path],
            ["catdoc", "-d", "utf-8", tmp_path],
            ["catdoc", tmp_path],
        ):
            try:
                result = subprocess.run(
                    cmd, capture_output=True, timeout=30, check=False
                )
            except FileNotFoundError:
                errors.append(f"{cmd[0]} 未安装")
                continue
            stdout = _decode_doc_output(result.stdout or b"").strip()
            stderr = (result.stderr or b"").decode("utf-8", errors="ignore").strip()
            if result.returncode == 0 and stdout:
                return stdout
            errors.append(
                f"{cmd[0]} 退出码={result.returncode} stderr={stderr[:120]}"
            )

        raise ValueError(
            "无法解析此 .doc 文件，可能为加密、超旧版本（Word 95 之前）或损坏文档。"
            "建议在 Word/WPS 中另存为 .docx 后再上传。"
            f"（详细：{'; '.join(errors)[:300]}）"
        )
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _extract_pseudo_doc_html(content: bytes) -> str | None:
    """Extract text from Confluence/Office HTML exported as `.doc`.

    Some systems export a MIME or quoted-printable HTML document but name it
    `.doc`. antiword can return headers or mojibake for this file type, so we
    detect it before invoking Word-specific extractors.
    """
    head = content[:4096].lower()
    looks_like_html_doc = (
        b"<html" in head
        or b"content-type: multipart/" in head
        or b"quoted-printable" in head
        or b"exported from confluence" in head
    )
    if not looks_like_html_doc:
        return None

    html_bytes: bytes | None = None
    try:
        msg = BytesParser(policy=policy.default).parsebytes(content)
        if msg.is_multipart():
            for part in msg.walk():
                ctype = part.get_content_type()
                if ctype in {"text/html", "application/xhtml+xml"}:
                    payload = part.get_payload(decode=True)
                    if payload:
                        html_bytes = payload
                        break
        elif msg.get_content_type() in {"text/html", "application/xhtml+xml"}:
            html_bytes = msg.get_payload(decode=True)
    except Exception:  # noqa: BLE001
        html_bytes = None

    if not html_bytes:
        marker = content.lower().find(b"<html")
        if marker >= 0:
            html_bytes = quopri.decodestring(content[marker:])
        else:
            html_bytes = quopri.decodestring(content)

    text = _html_to_text(_decode_doc_output(html_bytes))
    return text if text.strip() else None


def _html_to_text(html: str) -> str:
    html = re.sub(r"(?is)<(script|style|xml)[^>]*>.*?</\1>", " ", html)
    html = re.sub(r"(?i)<br\s*/?>", "\n", html)
    html = re.sub(r"(?i)</(p|div|tr|h[1-6]|li|table)>", "\n", html)
    html = re.sub(r"(?i)<td[^>]*>", " | ", html)
    text = re.sub(r"(?s)<[^>]+>", " ", html)
    text = unescape(text)
    lines = [re.sub(r"[ \t\r\f\v]+", " ", line).strip() for line in text.splitlines()]
    compact: list[str] = []
    for line in lines:
        if line and (not compact or compact[-1] != line):
            compact.append(line)
    return "\n".join(compact)


def _decode_doc_output(data: bytes) -> str:
    """Decode extractor output with a Chinese-friendly quality heuristic.

    antiword/catdoc may emit UTF-8, GBK/GB18030, or mixed legacy bytes depending
    on the document and mapping table. Decoding only as UTF-8 can produce
    "successful" but unreadable mojibake. We try common encodings and choose the
    result with the best Chinese/readability score.
    """
    if not data:
        return ""

    candidates: list[tuple[str, str]] = []
    for enc in ("utf-8", "gb18030", "gbk", "big5", "cp1252"):
        try:
            candidates.append((enc, data.decode(enc, errors="strict")))
        except UnicodeDecodeError:
            candidates.append((enc, data.decode(enc, errors="replace")))

    best_enc, best_text = max(candidates, key=lambda item: _text_quality(item[1]))
    logger.debug("Decoded .doc extractor output using %s", best_enc)
    return best_text


def _text_quality(text: str) -> float:
    if not text:
        return -9999.0
    total = len(text)
    replacement = text.count("\ufffd")
    chinese = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    private_use = sum(1 for ch in text if "\ue000" <= ch <= "\uf8ff")
    printable = sum(1 for ch in text if ch.isprintable() or ch in "\n\r\t")
    mojibake_markers = sum(
        text.count(token)
        for token in (
            "Ã",
            "Â",
            "Ð",
            "Ñ",
            "�",
            "瀹",
            "鎵",
            "鐑",
            "绾",
            "鍗",
            "妯",
            "夎",
            "卞",
            "垚",
            "殑",
        )
    )
    # Chinese documents are the high-frequency case here, but keep printable text
    # usable for English docs too.
    return (
        (chinese * 4)
        + (printable / max(total, 1))
        - (replacement * 10)
        - (private_use * 8)
        - (mojibake_markers * 6)
    )
